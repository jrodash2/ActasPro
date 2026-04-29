from io import BytesIO
from unicodedata import normalize


def get_acta_export_content(acta):
    contenido_final = (acta.contenido_final or "").strip()
    if contenido_final:
        return contenido_final
    contenido_borrador = (acta.contenido_borrador or "").strip()
    return contenido_borrador


def build_acta_docx_bytes(acta, institucion_nombre="Sistema de Actas Consistoriales"):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    contenido = get_acta_export_content(acta)
    if not contenido:
        raise ValueError("El acta no tiene contenido para exportar.")

    document = Document()
    encabezado = document.add_paragraph(institucion_nombre)
    encabezado.alignment = WD_ALIGN_PARAGRAPH.CENTER
    encabezado.runs[0].font.size = Pt(11)

    titulo = document.add_paragraph(f"ACTA NÚMERO {acta.numero_acta}/{acta.anio}")
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_run = titulo.runs[0]
    titulo_run.bold = True
    titulo_run.font.size = Pt(14)

    document.add_paragraph("")
    for bloque in contenido.replace("\r\n", "\n").split("\n"):
        texto = bloque.strip()
        if not texto:
            document.add_paragraph("")
            continue
        parrafo = document.add_paragraph(texto)
        parrafo.paragraph_format.space_after = Pt(8)

    if "firma" not in contenido.lower():
        document.add_paragraph("")
        document.add_paragraph("Firmas:")
        document.add_paragraph("")
        document.add_paragraph("__________________________")
        document.add_paragraph("Moderador")
        document.add_paragraph("")
        document.add_paragraph("__________________________")
        document.add_paragraph("Secretario")

    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def build_acta_docx(acta, institucion_nombre="Sistema de Actas Consistoriales"):
    bytes_content = build_acta_docx_bytes(acta, institucion_nombre=institucion_nombre)
    stream = BytesIO(bytes_content)
    stream.seek(0)
    filename = f"acta-{acta.numero_acta}-{acta.anio}"
    return stream, _sanitize_filename(filename)


def _sanitize_filename(value):
    normalized = normalize("NFKD", value).encode("ascii", "ignore").decode("ascii")
    return "".join(char if char.isalnum() or char in "-_" else "-" for char in normalized).strip("-_") or "acta"
